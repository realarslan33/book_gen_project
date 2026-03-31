import os
from typing import Dict, List
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def compile_txt(book: Dict, chapters: List[Dict], output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, f"{book['title']}.txt")
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(f"{book['title']}\n\n")
        f.write(f"Outline:\n{book.get('outline', '')}\n\n")
        for chapter in chapters:
            f.write(f"Chapter {chapter['chapter_number']}: {chapter.get('chapter_title', '')}\n")
            f.write(f"{chapter.get('chapter_text', '')}\n\n")
    return file_path


def compile_docx(book: Dict, chapters: List[Dict], output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, f"{book['title']}.docx")
    doc = Document()
    doc.add_heading(book['title'], 0)
    doc.add_paragraph("Outline")
    doc.add_paragraph(book.get('outline', ''))
    for chapter in chapters:
        doc.add_heading(f"Chapter {chapter['chapter_number']}: {chapter.get('chapter_title', '')}", level=1)
        doc.add_paragraph(chapter.get('chapter_text', ''))
    doc.save(file_path)
    return file_path


def compile_pdf(book: Dict, chapters: List[Dict], output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, f"{book['title']}.pdf")
    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, book['title'])
    y -= 30
    c.setFont("Helvetica", 10)
    for line in (book.get('outline', '') or 'Outline not available').splitlines():
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:100])
        y -= 14
    for chapter in chapters:
        if y < 80:
            c.showPage()
            y = height - 50
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, f"Chapter {chapter['chapter_number']}: {chapter.get('chapter_title', '')}")
        y -= 20
        c.setFont("Helvetica", 10)
        for line in (chapter.get('chapter_text', '') or '').splitlines():
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(50, y, line[:100])
            y -= 14
        y -= 10
    c.save()
    return file_path
